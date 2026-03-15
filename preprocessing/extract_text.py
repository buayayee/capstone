"""
extract_text.py
---------------
Extracts raw text from either a PDF or a Word (.docx) document.
Scanned PDFs (no text layer) fall back to Tesseract OCR automatically.

Cross-platform: works on Windows and macOS.
Tesseract and Poppler paths are detected automatically per OS.

Usage:
    python extract_text.py --input path/to/document.pdf
    python extract_text.py --input path/to/instructions.docx
"""

import argparse
import logging
import os
import platform
import re
import warnings

from PIL import Image
import numpy as np
import pdfplumber
from docx import Document as DocxDocument

# Suppress noisy RapidOCR INFO logs
logging.getLogger("RapidOCR").setLevel(logging.WARNING)
warnings.filterwarnings("ignore")

# Lazy-loaded RapidOCR singleton — loads ONNX models once on first use
_rapidocr_instance = None

def _get_rapidocr():
    global _rapidocr_instance
    if _rapidocr_instance is None:
        from rapidocr import RapidOCR
        _rapidocr_instance = RapidOCR()
    return _rapidocr_instance


# (Tesseract and Poppler are no longer required — RapidOCR handles all image/scanned OCR)


# ── Text helpers ─────────────────────────────────────────────────────────────

def clean_text(text: str) -> str:
    """Normalize whitespace and strip non-printable characters."""
    text = re.sub(r"\s+", " ", text)
    text = re.sub(r"[^\x20-\x7E\n]", "", text)
    return text.strip()


# ── Extractors ───────────────────────────────────────────────────────────────

def extract_from_docx(filepath: str) -> str:
    """Extract all paragraph text from a Word (.docx) document."""
    doc = DocxDocument(filepath)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also extract text from tables inside the document
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())
    return clean_text("\n".join(paragraphs))


def extract_from_digital_pdf(filepath: str) -> str:
    """Extract text from a digital (text-layer) PDF using pdfplumber."""
    text = ""
    with pdfplumber.open(filepath) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return clean_text(text)


def extract_from_scanned_pdf(filepath: str) -> str:
    """
    Fallback OCR path for scanned PDFs (no embedded text layer).
    Converts each page to an image using PyMuPDF (fitz), then runs RapidOCR.
    No external tools required — all handled by pip packages.
    """
    print(f"  [OCR Fallback] Converting scanned PDF pages: {filepath}")
    try:
        import fitz  # PyMuPDF
    except ImportError:
        raise RuntimeError("PyMuPDF not installed. Run: pip install pymupdf")

    ocr = _get_rapidocr()
    doc = fitz.open(filepath)
    text = ""
    for i, page in enumerate(doc):
        pix = page.get_pixmap(dpi=200)
        img = np.frombuffer(pix.samples, dtype=np.uint8).reshape(pix.height, pix.width, pix.n)
        if pix.n == 4:  # RGBA -> RGB
            img = img[:, :, :3]
        result = ocr(img)
        page_text = "\n".join(result.txts) if result.txts else ""
        text += page_text + "\n"
        print(f"  [OCR] Page {i+1}/{len(doc)} done.")
    doc.close()
    return clean_text(text)


# ── Public interface ──────────────────────────────────────────────────────────

IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".bmp", ".tiff", ".tif", ".webp"}


def extract_from_image(filepath: str) -> str:
    """
    Run RapidOCR on an image file (JPG, PNG, BMP, TIFF, etc.).
    Uses the ONNX-based RapidOCR engine — no external tools required.
    Large images (>25 MP) are downsampled to ~25 MP before OCR to avoid
    PIL decompression-bomb limits while preserving OCR quality.
    """
    print(f"  [OCR Image] Running OCR on: {filepath}")
    ocr = _get_rapidocr()

    # Disable PIL decompression-bomb guard (we trust our own submission files)
    # and resize if the image is unreasonably large (>25 MP) to keep OCR fast.
    Image.MAX_IMAGE_PIXELS = None
    img_pil = Image.open(filepath).convert("RGB")
    w, h = img_pil.size
    _MAX_PIXELS = 25_000_000
    if w * h > _MAX_PIXELS:
        scale = (_MAX_PIXELS / (w * h)) ** 0.5
        new_w, new_h = max(1, int(w * scale)), max(1, int(h * scale))
        print(f"  [OCR Image] Resizing {w}x{h} -> {new_w}x{new_h} before OCR")
        img_pil = img_pil.resize((new_w, new_h), Image.LANCZOS)

    img = np.array(img_pil)
    result = ocr(img)
    text = "\n".join(result.txts) if result.txts else ""
    return clean_text(text)


def extract_text(filepath: str) -> str:
    """
    Extract text from a PDF (.pdf), Word document (.docx), or image file
    (.jpg, .jpeg, .png, .bmp, .tiff, .tif, .webp).
    PDF files automatically fall back to OCR if no text layer is detected.
    """
    if not os.path.exists(filepath):
        raise FileNotFoundError(f"File not found: {filepath}")

    ext = os.path.splitext(filepath)[1].lower()

    if ext == ".docx":
        return extract_from_docx(filepath)
    elif ext == ".pdf":
        text = extract_from_digital_pdf(filepath)
        if len(text.split()) < 20:  # scanned PDF — use OCR
            text = extract_from_scanned_pdf(filepath)
        return text
    elif ext in IMAGE_EXTENSIONS:
        return extract_from_image(filepath)
    else:
        raise ValueError(
            f"Unsupported file type '{ext}'. "
            f"Supported: .pdf, .docx, .jpg, .jpeg, .png, .bmp, .tiff, .tif, .webp"
        )


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Extract text from a PDF or Word document.")
    parser.add_argument("--input", required=True, help="Path to the .pdf or .docx file")
    args = parser.parse_args()

    result = extract_text(args.input)
    print("\n=== Extracted Text ===")
    print(result)
